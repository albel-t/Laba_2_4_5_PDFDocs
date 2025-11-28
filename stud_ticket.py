import re
from docx import Document
from openpyxl import load_workbook

def fill_template_from_excel(template_path, output_path, excel_path, data_row):
    """
    Заполняет шаблон Word данными из строки Excel.

    :param template_path: Путь к файлу-шаблону (.docx)
    :param output_path: Путь для сохранения заполненного документа
    :param excel_path: Путь к файлу Excel (.xlsx)
    :param data_row: Номер строки в Excel для обработки (начиная с 1)
    """
    # 1. Загружаем документы
    doc = Document(template_path)
    wb = load_workbook(excel_path)
    ws = wb.active

    # 2. Подготовим данные из Excel для строки
    excel_data = {}
    for col in ws.iter_cols(min_row=data_row, max_row=data_row, min_col=1, max_col=ws.max_column):
        col_letter = col[0].column_letter  # Получаем букву столбца (A, B, C...)
        cell_value = str(col[0].value) if col[0].value is not None else ""
        excel_data[col_letter] = cell_value

    # 3. Функция для форматирования текста до нужной длины
    def format_text(text, max_length):
        """Обрезает текст до max_length или добавляет пробелы в конце."""
        if len(text) > max_length:
            return text[:max_length]
        else:
            return text + ' ' * (max_length - len(text))

    # 4. Специальная обработка для столбца A (разбиение на части)
    if 'A' in excel_data:
        a_data = excel_data['A']
        # Временное разбиение, будет уточнено при обработке конкретных плейсхолдеров
        a_parts = [a_data[i:i+50] for i in range(0, len(a_data), 50)]
        # Ограничиваем тремя частями
        excel_data['A_PARTS'] = a_parts[:3]
    else:
        excel_data['A_PARTS'] = []

    # 5. Функция для замены текста в параграфах и таблицах с сохранением форматирования
    def replace_in_paragraph(paragraph, placeholder, new_text):
        if placeholder not in paragraph.text:
            return
        
        # Находим run, содержащий плейсхолдер
        for run in paragraph.runs:
            if placeholder in run.text:
                # Сохраняем форматирование этого run
                font_size = run.font.size
                font_name = run.font.name
                is_bold = run.font.bold
                is_italic = run.font.italic
                
                # Заменяем текст
                run.text = run.text.replace(placeholder, new_text)
                
                # Восстанавливаем форматирование
                run.font.size = font_size
                run.font.name = font_name
                run.font.bold = is_bold
                run.font.italic = is_italic
                return
        
        # Если плейсхолдер разбит на несколько runs, используем комплексную замену
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder in full_text:
            # Находим первый run, который содержит часть плейсхолдера
            format_run = None
            for run in paragraph.runs:
                if any(char in run.text for char in placeholder):
                    format_run = run
                    break
            
            if format_run:
                # Полностью перестраиваем параграф
                new_full_text = full_text.replace(placeholder, new_text)
                
                # Очищаем все runs
                for run in paragraph.runs:
                    run.text = ""
                
                # Создаем новый run с сохраненным форматированием
                new_run = paragraph.add_run(new_full_text)
                new_run.font.size = format_run.font.size
                new_run.font.name = format_run.font.name
                new_run.font.bold = format_run.font.bold
                new_run.font.italic = format_run.font.italic

    # 6. Поиск и замена плейсхолдеров в документе
    # Регулярное выражение для поиска плейсхолдеров в формате -X-C- и -aN-C-
    pattern = r'-([A-Z]\d?)-(\d+)-'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        for match in matches:
            col_key, max_chars = match
            max_chars = int(max_chars)
            placeholder = f"-{col_key}-{max_chars}-"
            
            new_text = ""
            if col_key.startswith('A'):  # Обработка специальных слотов для столбца A
                part_num = int(col_key[1])
                if 'A_PARTS' in excel_data and len(excel_data['A_PARTS']) >= part_num:
                    part_text = excel_data['A_PARTS'][part_num-1]
                    new_text = format_text(part_text, max_chars)  # Обрезаем или добавляем пробелы
            else:  # Обычные столбцы (B, C, D...)
                if col_key in excel_data:
                    new_text = format_text(excel_data[col_key], max_chars)  # Обрезаем или добавляем пробелы
            
            if new_text:
                replace_in_paragraph(paragraph, placeholder, new_text)

    # 7. Также проверяем таблицы в документе
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    for match in matches:
                        col_key, max_chars = match
                        max_chars = int(max_chars)
                        placeholder = f"-{col_key}-{max_chars}-"
                        
                        new_text = ""
                        if col_key.startswith('a'):
                            part_num = int(col_key[1])
                            if 'A_PARTS' in excel_data and len(excel_data['A_PARTS']) >= part_num:
                                part_text = excel_data['A_PARTS'][part_num-1]
                                new_text = format_text(part_text, max_chars)
                        else:
                            if col_key in excel_data:
                                new_text = format_text(excel_data[col_key], max_chars)
                        
                        if new_text:
                            replace_in_paragraph(paragraph, placeholder, new_text)

    # 8. Сохраняем заполненный документ
    doc.save(output_path)
    print(f"Документ сохранен как: {output_path}")

# Пример использования
if __name__ == "__main__":
    fill_template_from_excel(
        template_path="D:/шаблон_студенческий.docx",
        output_path="заполненный_документ.docx", 
        excel_path="D:/данные.xlsx",
        data_row=3  # Обрабатываем строку 3 (где первая строка с данными)
    )